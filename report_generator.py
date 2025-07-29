# report_generator.py - Automated report generation system

import os
import json
import pandas as pd
import numpy as np
from datetime import datetime, timedelta
from pathlib import Path
import matplotlib.pyplot as plt
import seaborn as sns
from typing import Dict, List, Optional, Any, Tuple
import logging
from dataclasses import dataclass
from jinja2 import Template, Environment, FileSystemLoader
import pdfkit
import plotly.graph_objects as go
import plotly.express as px
from plotly.subplots import make_subplots
import sqlite3
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

@dataclass
class ReportConfig:
    """Configuration for report generation"""
    title: str
    subtitle: str
    report_type: str  # 'executive', 'technical', 'comparative', 'custom'
    format: str  # 'html', 'pdf', 'docx', 'markdown'
    include_visualizations: bool = True
    include_raw_data: bool = False
    language: str = 'en'
    branding: Dict[str, Any] = None
    recipients: List[str] = None

class ReportGenerator:
    """Comprehensive automated report generation system"""
    
    def __init__(self, 
                 data_path: str = "./data",
                 output_dir: str = "./reports",
                 template_dir: str = "./templates"):
        self.data_path = Path(data_path)
        self.output_dir = Path(output_dir)
        self.template_dir = Path(template_dir)
        
        # Create directories
        self.output_dir.mkdir(exist_ok=True)
        self.template_dir.mkdir(exist_ok=True)
        
        # Initialize template environment
        self.env = Environment(loader=FileSystemLoader(str(self.template_dir)))
        
        # Report sections
        self.sections = {
            'executive': ['summary', 'key_findings', 'recommendations', 'metrics_overview'],
            'technical': ['methodology', 'detailed_analysis', 'statistical_tests', 'technical_metrics'],
            'comparative': ['comparison_overview', 'model_analysis', 'language_analysis', 'trends'],
            'custom': []  # Flexible sections
        }
        
        # Create default templates if they don't exist
        self._create_default_templates()
    
    def generate_report(self,
                       dataset_id: Optional[str] = None,
                       date_range: Optional[Tuple[datetime, datetime]] = None,
                       config: Optional[ReportConfig] = None) -> str:
        """Generate a comprehensive report"""
        
        if not config:
            config = ReportConfig(
                title="MASB-Alt Evaluation Report",
                subtitle=f"Generated on {datetime.now().strftime('%Y-%m-%d')}",
                report_type="executive",
                format="html"
            )
        
        logger.info(f"Generating {config.report_type} report in {config.format} format")
        
        # Load data
        data = self._load_report_data(dataset_id, date_range)
        
        if data['evaluations'].empty:
            logger.warning("No data available for report generation")
            return None
        
        # Prepare report content
        report_content = self._prepare_report_content(data, config)
        
        # Generate visualizations if needed
        if config.include_visualizations:
            report_content['visualizations'] = self._generate_visualizations(data, config)
        
        # Generate report in requested format
        output_file = self._generate_output(report_content, config)
        
        # Send report if recipients specified
        if config.recipients:
            self._send_report(output_file, config)
        
        logger.info(f"Report generated: {output_file}")
        return output_file
    
    def _load_report_data(self, 
                         dataset_id: Optional[str] = None,
                         date_range: Optional[Tuple[datetime, datetime]] = None) -> Dict[str, Any]:
        """Load all necessary data for report generation"""
        db_path = self.data_path / "masb_alt.db"
        conn = sqlite3.connect(db_path)
        
        # Base query
        base_query = """
        SELECT 
            e.*,
            r.model,
            r.latency_ms,
            r.token_count,
            r.error as response_error,
            p.language,
            p.domain,
            p.text as prompt_text,
            p.risk_level as prompt_risk_level,
            d.name as dataset_name
        FROM evaluations e
        JOIN responses r ON e.response_id = r.response_id
        JOIN prompts p ON r.prompt_id = p.prompt_id
        LEFT JOIN datasets d ON p.dataset_id = d.dataset_id
        WHERE 1=1
        """
        
        params = []
        
        if dataset_id:
            base_query += " AND p.dataset_id = ?"
            params.append(dataset_id)
        
        if date_range:
            base_query += " AND e.timestamp BETWEEN ? AND ?"
            params.extend([date_range[0].isoformat(), date_range[1].isoformat()])
        
        # Load main evaluation data
        evaluations_df = pd.read_sql_query(base_query, conn, params=params)
        evaluations_df['timestamp'] = pd.to_datetime(evaluations_df['timestamp'])
        evaluations_df['risk_flags'] = evaluations_df['risk_flags'].apply(
            lambda x: json.loads(x) if x else {}
        )
        
        # Load summary statistics
        stats_query = """
        SELECT * FROM corpus_stats
        ORDER BY last_updated DESC
        """
        stats_df = pd.read_sql_query(stats_query, conn)
        
        conn.close()
        
        return {
            'evaluations': evaluations_df,
            'statistics': stats_df,
            'metadata': {
                'total_evaluations': len(evaluations_df),
                'date_range': [
                    evaluations_df['timestamp'].min() if not evaluations_df.empty else None,
                    evaluations_df['timestamp'].max() if not evaluations_df.empty else None
                ],
                'models': evaluations_df['model'].unique().tolist() if not evaluations_df.empty else [],
                'languages': evaluations_df['language'].unique().tolist() if not evaluations_df.empty else [],
                'domains': evaluations_df['domain'].unique().tolist() if not evaluations_df.empty else []
            }
        }
    
    def _prepare_report_content(self, data: Dict[str, Any], config: ReportConfig) -> Dict[str, Any]:
        """Prepare content for report based on type"""
        content = {
            'title': config.title,
            'subtitle': config.subtitle,
            'generated_at': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
            'metadata': data['metadata']
        }
        
        df = data['evaluations']
        
        if config.report_type == 'executive':
            content.update(self._prepare_executive_content(df))
        elif config.report_type == 'technical':
            content.update(self._prepare_technical_content(df))
        elif config.report_type == 'comparative':
            content.update(self._prepare_comparative_content(df))
        else:  # custom
            content.update(self._prepare_custom_content(df, config))
        
        return content
    
    def _prepare_executive_content(self, df: pd.DataFrame) -> Dict[str, Any]:
        """Prepare executive summary content"""
        content = {}
        
        # Summary metrics
        content['summary'] = {
            'total_evaluations': len(df),
            'average_alignment_score': df['alignment_score'].mean(),
            'high_risk_percentage': (df['risk_level'].isin(['high', 'critical'])).mean() * 100,
            'evaluation_period': f"{df['timestamp'].min().date()} to {df['timestamp'].max().date()}",
            'models_evaluated': df['model'].nunique(),
            'languages_covered': df['language'].nunique(),
            'domains_assessed': df['domain'].nunique()
        }
        
        # Key findings
        findings = []
        
        # Best performing model
        model_scores = df.groupby('model')['alignment_score'].mean().sort_values(ascending=False)
        if len(model_scores) > 0:
            findings.append({
                'type': 'positive',
                'title': 'Best Performing Model',
                'detail': f"{model_scores.index[0]} achieved the highest average alignment score of {model_scores.iloc[0]:.2f}"
            })
        
        # Language gaps
        lang_scores = df.groupby('language')['alignment_score'].mean()
        lang_gap = lang_scores.max() - lang_scores.min()
        if lang_gap > 0.5:
            findings.append({
                'type': 'concern',
                'title': 'Language Performance Gap',
                'detail': f"Significant gap of {lang_gap:.2f} points between best and worst performing languages"
            })
        
        # Risk areas
        risk_by_domain = df.groupby('domain')['risk_level'].apply(
            lambda x: (x.isin(['high', 'critical'])).mean() * 100
        ).sort_values(ascending=False)
        
        if risk_by_domain.iloc[0] > 20:
            findings.append({
                'type': 'warning',
                'title': 'High Risk Domain',
                'detail': f"{risk_by_domain.index[0]} domain shows {risk_by_domain.iloc[0]:.1f}% high-risk responses"
            })
        
        content['key_findings'] = findings
        
        # Recommendations
        recommendations = []
        
        # Model recommendations
        if model_scores.iloc[-1] < 3.0:
            recommendations.append({
                'priority': 'high',
                'area': 'Model Performance',
                'recommendation': f"Consider replacing or fine-tuning {model_scores.index[-1]} due to consistently low scores"
            })
        
        # Language recommendations
        low_performing_langs = lang_scores[lang_scores < 3.5].index.tolist()
        if low_performing_langs:
            recommendations.append({
                'priority': 'medium',
                'area': 'Language Support',
                'recommendation': f"Improve support for: {', '.join(low_performing_langs)}"
            })
        
        content['recommendations'] = recommendations
        
        # Metrics overview
        content['metrics_overview'] = {
            'by_model': df.groupby('model').agg({
                'alignment_score': ['mean', 'std', 'count'],
                'latency_ms': 'mean'
            }).to_dict(),
            'by_language': df.groupby('language').agg({
                'alignment_score': ['mean', 'count']
            }).to_dict(),
            'by_domain': df.groupby('domain').agg({
                'alignment_score': 'mean',
                'risk_level': lambda x: (x.isin(['high', 'critical'])).mean() * 100
            }).to_dict()
        }
        
        return content
    
    def _prepare_technical_content(self, df: pd.DataFrame) -> Dict[str, Any]:
        """Prepare technical report content"""
        content = {}
        
        # Methodology
        content['methodology'] = {
            'evaluation_framework': 'MASB-Alt Multilingual Alignment Safety Benchmark',
            'scoring_system': '1-5 scale alignment scoring with risk flag detection',
            'languages_tested': df['language'].unique().tolist(),
            'domains_covered': df['domain'].unique().tolist(),
            'models_evaluated': df['model'].unique().tolist(),
            'total_prompts': df['prompt_text'].nunique(),
            'evaluation_period': f"{df['timestamp'].min()} to {df['timestamp'].max()}"
        }
        
        # Detailed analysis
        detailed_analysis = {}
        
        # Performance metrics by model
        for model in df['model'].unique():
            model_data = df[df['model'] == model]
            
            # Calculate percentiles
            percentiles = model_data['alignment_score'].quantile([0.25, 0.5, 0.75]).to_dict()
            
            # Risk flag analysis
            risk_flags = {}
            for _, row in model_data.iterrows():
                for flag, value in row['risk_flags'].items():
                    if value:
                        risk_flags[flag] = risk_flags.get(flag, 0) + 1
            
            detailed_analysis[model] = {
                'n_evaluations': len(model_data),
                'mean_score': model_data['alignment_score'].mean(),
                'std_score': model_data['alignment_score'].std(),
                'percentiles': percentiles,
                'mean_latency_ms': model_data['latency_ms'].mean(),
                'error_rate': (model_data['response_error'].notna()).mean() * 100,
                'risk_flags_frequency': risk_flags,
                'confidence_score': model_data['confidence_score'].mean()
            }
        
        content['detailed_analysis'] = detailed_analysis
        
        # Statistical tests
        from scipy import stats
        
        statistical_tests = {}
        
        # ANOVA for model differences
        model_groups = [df[df['model'] == m]['alignment_score'].values 
                       for m in df['model'].unique()]
        if len(model_groups) > 1:
            f_stat, p_value = stats.f_oneway(*model_groups)
            statistical_tests['model_anova'] = {
                'test': 'One-way ANOVA',
                'f_statistic': f_stat,
                'p_value': p_value,
                'significant': p_value < 0.05
            }
        
        # Correlation analysis
        if len(df) > 30:
            corr, p_val = stats.pearsonr(df['latency_ms'], df['alignment_score'])
            statistical_tests['latency_score_correlation'] = {
                'test': 'Pearson Correlation',
                'correlation': corr,
                'p_value': p_val,
                'interpretation': 'negative correlation' if corr < -0.3 else 'no significant correlation'
            }
        
        content['statistical_tests'] = statistical_tests
        
        # Technical metrics
        content['technical_metrics'] = {
            'response_time_analysis': {
                'mean_latency_ms': df['latency_ms'].mean(),
                'median_latency_ms': df['latency_ms'].median(),
                'p95_latency_ms': df['latency_ms'].quantile(0.95),
                'p99_latency_ms': df['latency_ms'].quantile(0.99)
            },
            'token_usage': {
                'mean_tokens': df['token_count'].mean() if 'token_count' in df.columns else None,
                'total_tokens': df['token_count'].sum() if 'token_count' in df.columns else None
            },
            'error_analysis': {
                'total_errors': df['response_error'].notna().sum(),
                'error_rate': (df['response_error'].notna()).mean() * 100,
                'errors_by_model': df.groupby('model')['response_error'].apply(lambda x: x.notna().sum()).to_dict()
            }
        }
        
        return content
    
    def _prepare_comparative_content(self, df: pd.DataFrame) -> Dict[str, Any]:
        """Prepare comparative analysis content"""
        content = {}
        
        # Comparison overview
        content['comparison_overview'] = {
            'models_compared': df['model'].unique().tolist(),
            'comparison_period': f"{df['timestamp'].min().date()} to {df['timestamp'].max().date()}",
            'total_evaluations': len(df),
            'evaluation_distribution': df['model'].value_counts().to_dict()
        }
        
        # Model analysis
        model_comparison = pd.DataFrame()
        for model in df['model'].unique():
            model_data = df[df['model'] == model]
            model_comparison = pd.concat([model_comparison, pd.DataFrame({
                'model': [model],
                'avg_score': [model_data['alignment_score'].mean()],
                'std_score': [model_data['alignment_score'].std()],
                'high_risk_rate': [(model_data['risk_level'].isin(['high', 'critical'])).mean() * 100],
                'avg_latency': [model_data['latency_ms'].mean()],
                'sample_size': [len(model_data)]
            })])
        
        content['model_analysis'] = model_comparison.to_dict('records')
        
        # Language analysis
        language_comparison = {}
        for lang in df['language'].unique():
            lang_data = df[df['language'] == lang]
            language_comparison[lang] = {
                'models': {},
                'overall_score': lang_data['alignment_score'].mean()
            }
            
            for model in df['model'].unique():
                model_lang_data = lang_data[lang_data['model'] == model]
                if len(model_lang_data) > 0:
                    language_comparison[lang]['models'][model] = {
                        'avg_score': model_lang_data['alignment_score'].mean(),
                        'n_samples': len(model_lang_data)
                    }
        
        content['language_analysis'] = language_comparison
        
        # Trends
        df['date'] = df['timestamp'].dt.date
        daily_trends = df.groupby(['date', 'model'])['alignment_score'].mean().reset_index()
        
        content['trends'] = {
            'daily_scores': daily_trends.to_dict('records'),
            'trend_summary': 'Scores show stability' if df['alignment_score'].std() < 0.5 else 'Significant variability observed'
        }
        
        return content
    
    def _prepare_custom_content(self, df: pd.DataFrame, config: ReportConfig) -> Dict[str, Any]:
        """Prepare custom report content based on configuration"""
        content = {}
        
        # Add basic statistics
        content['basic_stats'] = {
            'total_evaluations': len(df),
            'models': df['model'].unique().tolist(),
            'languages': df['language'].unique().tolist(),
            'domains': df['domain'].unique().tolist()
        }
        
        # Add any custom sections specified in config
        if hasattr(config, 'custom_sections'):
            for section in config.custom_sections:
                if section == 'risk_analysis':
                    content['risk_analysis'] = self._analyze_risks(df)
                elif section == 'performance_matrix':
                    content['performance_matrix'] = self._create_performance_matrix(df)
                # Add more custom sections as needed
        
        return content
    
    def _analyze_risks(self, df: pd.DataFrame) -> Dict[str, Any]:
        """Detailed risk analysis"""
        risk_analysis = {
            'by_type': {},
            'by_domain': {},
            'high_risk_examples': []
        }
        
        # Count risk flags
        risk_counts = {}
        for _, row in df.iterrows():
            for flag, value in row['risk_flags'].items():
                if value:
                    risk_counts[flag] = risk_counts.get(flag, 0) + 1
        
        risk_analysis['by_type'] = risk_counts
        
        # Risk by domain
        for domain in df['domain'].unique():
            domain_data = df[df['domain'] == domain]
            risk_analysis['by_domain'][domain] = {
                'high_risk_rate': (domain_data['risk_level'].isin(['high', 'critical'])).mean() * 100,
                'total_evaluations': len(domain_data)
            }
        
        # High risk examples
        high_risk = df[df['risk_level'].isin(['high', 'critical'])].head(5)
        for _, row in high_risk.iterrows():
            risk_analysis['high_risk_examples'].append({
                'prompt': row['prompt_text'][:100] + '...',
                'model': row['model'],
                'risk_level': row['risk_level'],
                'flags': [k for k, v in row['risk_flags'].items() if v]
            })
        
        return risk_analysis
    
    def _create_performance_matrix(self, df: pd.DataFrame) -> List[Dict[str, Any]]:
        """Create performance matrix data"""
        matrix = []
        
        for model in df['model'].unique():
            for language in df['language'].unique():
                for domain in df['domain'].unique():
                    subset = df[(df['model'] == model) & 
                              (df['language'] == language) & 
                              (df['domain'] == domain)]
                    
                    if len(subset) > 0:
                        matrix.append({
                            'model': model,
                            'language': language,
                            'domain': domain,
                            'avg_score': subset['alignment_score'].mean(),
                            'n_samples': len(subset)
                        })
        
        return matrix
    
    def _generate_visualizations(self, data: Dict[str, Any], config: ReportConfig) -> Dict[str, str]:
        """Generate all visualizations for the report"""
        viz_files = {}
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        df = data['evaluations']
        
        # 1. Overall performance chart
        fig = plt.figure(figsize=(10, 6))
        model_scores = df.groupby('model')['alignment_score'].mean().sort_values()
        model_scores.plot(kind='barh', color='skyblue')
        plt.xlabel('Average Alignment Score')
        plt.title('Model Performance Comparison')
        plt.tight_layout()
        
        filename = self.output_dir / f"model_performance_{timestamp}.png"
        plt.savefig(filename, dpi=300, bbox_inches='tight')
        plt.close()
        viz_files['model_performance'] = str(filename)
        
        # 2. Language-Domain heatmap
        fig = plt.figure(figsize=(12, 8))
        pivot_table = df.pivot_table(
            values='alignment_score',
            index='language',
            columns='domain',
            aggfunc='mean'
        )
        sns.heatmap(pivot_table, annot=True, fmt='.2f', cmap='RdYlGn', center=3)
        plt.title('Performance Heatmap: Language vs Domain')
        plt.tight_layout()
        
        filename = self.output_dir / f"language_domain_heatmap_{timestamp}.png"
        plt.savefig(filename, dpi=300, bbox_inches='tight')
        plt.close()
        viz_files['language_domain_heatmap'] = str(filename)
        
        # 3. Risk distribution
        fig = plt.figure(figsize=(8, 8))
        risk_counts = df['risk_level'].value_counts()
        colors = {'low': 'green', 'medium': 'yellow', 'high': 'orange', 'critical': 'red'}
        plt.pie(risk_counts.values, labels=risk_counts.index, autopct='%1.1f%%',
               colors=[colors.get(x, 'gray') for x in risk_counts.index])
        plt.title('Risk Level Distribution')
        
        filename = self.output_dir / f"risk_distribution_{timestamp}.png"
        plt.savefig(filename, dpi=300, bbox_inches='tight')
        plt.close()
        viz_files['risk_distribution'] = str(filename)
        
        # 4. Temporal trends (if enough data)
        if len(df) > 50:
            fig = plt.figure(figsize=(12, 6))
            df_sorted = df.sort_values('timestamp')
            df_sorted['date'] = df_sorted['timestamp'].dt.date
            
            for model in df['model'].unique():
                model_data = df_sorted[df_sorted['model'] == model]
                daily_avg = model_data.groupby('date')['alignment_score'].mean()
                plt.plot(daily_avg.index, daily_avg.values, marker='o', label=model)
            
            plt.xlabel('Date')
            plt.ylabel('Average Alignment Score')
            plt.title('Alignment Score Trends Over Time')
            plt.legend()
            plt.xticks(rotation=45)
            plt.tight_layout()
            
            filename = self.output_dir / f"temporal_trends_{timestamp}.png"
            plt.savefig(filename, dpi=300, bbox_inches='tight')
            plt.close()
            viz_files['temporal_trends'] = str(filename)
        
        return viz_files
    
    def _generate_output(self, content: Dict[str, Any], config: ReportConfig) -> str:
        """Generate report in the specified format"""
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        
        if config.format == 'html':
            return self._generate_html_report(content, config, timestamp)
        elif config.format == 'pdf':
            return self._generate_pdf_report(content, config, timestamp)
        elif config.format == 'docx':
            return self._generate_docx_report(content, config, timestamp)
        elif config.format == 'markdown':
            return self._generate_markdown_report(content, config, timestamp)
        else:
            raise ValueError(f"Unsupported format: {config.format}")
    
    def _generate_html_report(self, content: Dict[str, Any], config: ReportConfig, timestamp: str) -> str:
        """Generate HTML report"""
        # Create or load template
        template_name = f"{config.report_type}_template.html"
        
        if not (self.template_dir / template_name).exists():
            template_content = self._get_default_html_template(config.report_type)
            with open(self.template_dir / template_name, 'w') as f:
                f.write(template_content)
        
        template = self.env.get_template(template_name)
        
        # Render template
        html_content = template.render(**content)
        
        # Save report
        output_file = self.output_dir / f"report_{config.report_type}_{timestamp}.html"
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write(html_content)
        
        return str(output_file)
    
    def _generate_pdf_report(self, content: Dict[str, Any], config: ReportConfig, timestamp: str) -> str:
        """Generate PDF report from HTML"""
        # First generate HTML
        html_file = self._generate_html_report(content, config, timestamp)
        
        # Convert to PDF
        pdf_file = html_file.replace('.html', '.pdf')
        
        try:
            pdfkit.from_file(html_file, pdf_file)
            # Optionally remove HTML file
            os.remove(html_file)
        except Exception as e:
            logger.error(f"PDF generation failed: {e}")
            return html_file
        
        return pdf_file
    
    def _generate_docx_report(self, content: Dict[str, Any], config: ReportConfig, timestamp: str) -> str:
        """Generate DOCX report"""
        doc = Document()
        
        # Add title
        title = doc.add_heading(content['title'], 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add subtitle
        subtitle = doc.add_heading(content['subtitle'], 1)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add generation info
        doc.add_paragraph(f"Generated on: {content['generated_at']}")
        doc.add_paragraph()
        
        # Add content based on report type
        if config.report_type == 'executive':
            self._add_executive_sections_docx(doc, content)
        elif config.report_type == 'technical':
            self._add_technical_sections_docx(doc, content)
        
        # Add visualizations
        if 'visualizations' in content:
            doc.add_page_break()
            doc.add_heading('Visualizations', 1)
            for viz_name, viz_path in content['visualizations'].items():
                if os.path.exists(viz_path):
                    doc.add_paragraph(viz_name.replace('_', ' ').title())
                    doc.add_picture(viz_path, width=Inches(6))
                    doc.add_paragraph()
        
        # Save document
        output_file = self.output_dir / f"report_{config.report_type}_{timestamp}.docx"
        doc.save(output_file)
        
        return str(output_file)
    
    def _add_executive_sections_docx(self, doc: Document, content: Dict[str, Any]):
        """Add executive report sections to DOCX"""
        # Summary
        doc.add_heading('Executive Summary', 1)
        for key, value in content['summary'].items():
            doc.add_paragraph(f"{key.replace('_', ' ').title()}: {value}")
        
        # Key Findings
        doc.add_heading('Key Findings', 1)
        for finding in content['key_findings']:
            p = doc.add_paragraph()
            if finding['type'] == 'positive':
                p.add_run('✓ ').font.color.rgb = RGBColor(0, 128, 0)
            elif finding['type'] == 'warning':
                p.add_run('⚠ ').font.color.rgb = RGBColor(255, 165, 0)
            else:
                p.add_run('✗ ').font.color.rgb = RGBColor(255, 0, 0)
            
            p.add_run(f"{finding['title']}: {finding['detail']}")
        
        # Recommendations
        doc.add_heading('Recommendations', 1)
        for rec in content['recommendations']:
            p = doc.add_paragraph()
            p.add_run(f"[{rec['priority'].upper()}] ").bold = True
            p.add_run(f"{rec['area']}: {rec['recommendation']}")
    
    def _add_technical_sections_docx(self, doc: Document, content: Dict[str, Any]):
        """Add technical report sections to DOCX"""
        # Methodology
        doc.add_heading('Methodology', 1)
        for key, value in content['methodology'].items():
            doc.add_paragraph(f"{key.replace('_', ' ').title()}: {value}")
        
        # Detailed Analysis
        doc.add_heading('Detailed Analysis', 1)
        for model, metrics in content['detailed_analysis'].items():
            doc.add_heading(model, 2)
            for metric, value in metrics.items():
                if isinstance(value, dict):
                    doc.add_paragraph(f"{metric.replace('_', ' ').title()}:")
                    for k, v in value.items():
                        doc.add_paragraph(f"  • {k}: {v}", style='List Bullet')
                else:
                    doc.add_paragraph(f"{metric.replace('_', ' ').title()}: {value}")
    
    def _generate_markdown_report(self, content: Dict[str, Any], config: ReportConfig, timestamp: str) -> str:
        """Generate Markdown report"""
        md_content = f"# {content['title']}\n\n"
        md_content += f"## {content['subtitle']}\n\n"
        md_content += f"*Generated on: {content['generated_at']}*\n\n"
        
        # Add content based on report type
        if config.report_type == 'executive':
            md_content += self._add_executive_sections_markdown(content)
        elif config.report_type == 'technical':
            md_content += self._add_technical_sections_markdown(content)
        
        # Save report
        output_file = self.output_dir / f"report_{config.report_type}_{timestamp}.md"
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write(md_content)
        
        return str(output_file)
    
    def _add_executive_sections_markdown(self, content: Dict[str, Any]) -> str:
        """Add executive sections in Markdown format"""
        md = "## Executive Summary\n\n"
        
        for key, value in content['summary'].items():
            md += f"- **{key.replace('_', ' ').title()}**: {value}\n"
        
        md += "\n## Key Findings\n\n"
        for finding in content['key_findings']:
            icon = "✅" if finding['type'] == 'positive' else "⚠️" if finding['type'] == 'warning' else "❌"
            md += f"{icon} **{finding['title']}**: {finding['detail']}\n\n"
        
        md += "## Recommendations\n\n"
        for rec in content['recommendations']:
            md += f"### [{rec['priority'].upper()}] {rec['area']}\n"
            md += f"{rec['recommendation']}\n\n"
        
        return md
    
    def _add_technical_sections_markdown(self, content: Dict[str, Any]) -> str:
        """Add technical sections in Markdown format"""
        md = "## Methodology\n\n"
        
        for key, value in content['methodology'].items():
            md += f"- **{key.replace('_', ' ').title()}**: {value}\n"
        
        md += "\n## Detailed Analysis\n\n"
        for model, metrics in content['detailed_analysis'].items():
            md += f"### {model}\n\n"
            for metric, value in metrics.items():
                if isinstance(value, dict):
                    md += f"**{metric.replace('_', ' ').title()}**:\n"
                    for k, v in value.items():
                        md += f"- {k}: {v}\n"
                    md += "\n"
                else:
                    md += f"- **{metric.replace('_', ' ').title()}**: {value}\n"
            md += "\n"
        
        return md
    
    def _send_report(self, report_file: str, config: ReportConfig):
        """Send report via email"""
        if not config.recipients:
            return
        
        # Email configuration (you'd need to set these)
        smtp_server = os.getenv('SMTP_SERVER', 'smtp.gmail.com')
        smtp_port = int(os.getenv('SMTP_PORT', '587'))
        smtp_username = os.getenv('SMTP_USERNAME')
        smtp_password = os.getenv('SMTP_PASSWORD')
        
        if not smtp_username or not smtp_password:
            logger.warning("Email credentials not configured")
            return
        
        try:
            # Create message
            msg = MIMEMultipart()
            msg['From'] = smtp_username
            msg['To'] = ', '.join(config.recipients)
            msg['Subject'] = f"MASB-Alt Report: {config.title}"
            
            # Add body
            body = f"""
            Dear Team,
            
            Please find attached the {config.report_type} report: {config.title}
            
            Report generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
            
            Best regards,
            MASB-Alt Reporting System
            """
            msg.attach(MIMEText(body, 'plain'))
            
            # Add attachment
            with open(report_file, 'rb') as f:
                part = MIMEBase('application', 'octet-stream')
                part.set_payload(f.read())
                encoders.encode_base64(part)
                part.add_header(
                    'Content-Disposition',
                    f'attachment; filename= {os.path.basename(report_file)}'
                )
                msg.attach(part)
            
            # Send email
            server = smtplib.SMTP(smtp_server, smtp_port)
            server.starttls()
            server.login(smtp_username, smtp_password)
            server.send_message(msg)
            server.quit()
            
            logger.info(f"Report sent to: {', '.join(config.recipients)}")
            
        except Exception as e:
            logger.error(f"Failed to send email: {e}")
    
    def _create_default_templates(self):
        """Create default report templates"""
        # Executive template
        executive_template = """
<!DOCTYPE html>
<html>
<head>
    <title>{{ title }}</title>
    <style>
        body { font-family: Arial, sans-serif; margin: 40px; line-height: 1.6; }
        h1, h2, h3 { color: #333; }
        .metric { background: #f4f4f4; padding: 15px; margin: 10px 0; border-radius: 5px; }
        .positive { color: green; }
        .warning { color: orange; }
        .negative { color: red; }
        table { border-collapse: collapse; width: 100%; margin: 20px 0; }
        th, td { border: 1px solid #ddd; padding: 8px; text-align: left; }
        th { background-color: #4CAF50; color: white; }
        .recommendation { background: #e8f4f8; padding: 15px; margin: 10px 0; border-left: 4px solid #2196F3; }
    </style>
</head>
<body>
    <h1>{{ title }}</h1>
    <h2>{{ subtitle }}</h2>
    <p><em>Generated on: {{ generated_at }}</em></p>
    
    <h2>Executive Summary</h2>
    <div class="metric">
        {% for key, value in summary.items() %}
        <p><strong>{{ key|replace('_', ' ')|title }}:</strong> {{ value }}</p>
        {% endfor %}
    </div>
    
    <h2>Key Findings</h2>
    {% for finding in key_findings %}
    <div class="{{ finding.type }}">
        <h3>{{ finding.title }}</h3>
        <p>{{ finding.detail }}</p>
    </div>
    {% endfor %}
    
    <h2>Recommendations</h2>
    {% for rec in recommendations %}
    <div class="recommendation">
        <h3>[{{ rec.priority|upper }}] {{ rec.area }}</h3>
        <p>{{ rec.recommendation }}</p>
    </div>
    {% endfor %}
    
    {% if visualizations %}
    <h2>Visualizations</h2>
    {% for name, path in visualizations.items() %}
    <div style="margin: 20px 0;">
        <h3>{{ name|replace('_', ' ')|title }}</h3>
        <img src="{{ path }}" style="max-width: 100%; height: auto;">
    </div>
    {% endfor %}
    {% endif %}
</body>
</html>
"""
        
        if not (self.template_dir / 'executive_template.html').exists():
            with open(self.template_dir / 'executive_template.html', 'w') as f:
                f.write(executive_template)
    
    def _get_default_html_template(self, report_type: str) -> str:
        """Get default HTML template for report type"""
        # Return the executive template as default
        return self._create_default_templates()

    def schedule_report(self, 
                       schedule: str,  # 'daily', 'weekly', 'monthly'
                       config: ReportConfig,
                       dataset_id: Optional[str] = None):
        """Schedule automatic report generation"""
        # This would integrate with a task scheduler like celery or APScheduler
        logger.info(f"Report scheduled: {schedule}")
        # Implementation would depend on your scheduling system

# CLI interface
def main():
    """Command-line interface for report generation"""
    import argparse
    
    parser = argparse.ArgumentParser(description="MASB-Alt Report Generator")
    parser.add_argument("--type", choices=['executive', 'technical', 'comparative', 'custom'],
                       default='executive', help="Report type")
    parser.add_argument("--format", choices=['html', 'pdf', 'docx', 'markdown'],
                       default='html', help="Output format")
    parser.add_argument("--dataset", help="Dataset ID")
    parser.add_argument("--days", type=int, default=7, help="Number of days to include")
    parser.add_argument("--email", nargs="+", help="Email recipients")
    parser.add_argument("--title", help="Report title")
    
    args = parser.parse_args()
    
    # Prepare configuration
    config = ReportConfig(
        title=args.title or f"MASB-Alt {args.type.title()} Report",
        subtitle=f"Analysis Period: Last {args.days} days",
        report_type=args.type,
        format=args.format,
        recipients=args.email
    )
    
    # Date range
    end_date = datetime.now()
    start_date = end_date - timedelta(days=args.days)
    
    # Generate report
    generator = ReportGenerator()
    report_file = generator.generate_report(
        dataset_id=args.dataset,
        date_range=(start_date, end_date),
        config=config
    )
    
    print(f"Report generated: {report_file}")
    
    if args.email:
        print(f"Report sent to: {', '.join(args.email)}")

if __name__ == "__main__":
    main()